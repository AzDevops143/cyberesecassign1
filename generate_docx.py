import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style definitions
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(51, 65, 85) # Slate 700

    # Custom Helper for Shading and Borders in Tables (Premium Look)
    def set_cell_background(cell, hex_color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(margin)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_cell_left_border(cell, hex_color, size="36"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), size)
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), hex_color)
        tcBorders.append(left)
        
        # Clear others
        for border_name in ['top', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # Add Document Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("IMPLEMENTATION GUIDE: HEXA FORCE LAB")
    run_title.font.name = 'Trebuchet MS'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Dark Slate 900
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("A 4-Stage Container Security Attack, Detection, and Defense Laboratory for DevSecOps Pipelines")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(14, 116, 144) # Cyan 700

    # Horizontal Divider Line
    divider_p = doc.add_paragraph()
    divider_p.paragraph_format.space_after = Pt(20)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12') # 1.5 pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0E7490') # Cyan 700
    pBdr.append(bottom)
    divider_p._p.get_or_add_pPr().append(pBdr)

    # Section Headers helper
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Trebuchet MS'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        
        # Add a subtle left accent border on paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24') # 3pt
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), '0E7490') # Cyan 700
        pBdr.append(left)
        pPr.append(pBdr)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Trebuchet MS'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(14, 116, 144) # Cyan 700
        return p

    # Code block helper (Styled single-cell table)
    def add_code_block(lines):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Lock width
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        set_cell_background(cell, "1F2937") # Gray 800 background
        set_cell_left_border(cell, "06B6D4", "24") # 3pt cyan left accent border
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        
        for i, line in enumerate(lines):
            if i > 0:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(243, 244, 246) # White-ish Gray 50
            
        doc.add_paragraph().paragraph_format.space_after = Pt(4) # spacer paragraph

    # Callout box helper
    def add_callout(text, title="NOTE", color_hex="0E7490", bg_hex="F0F9FF"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        set_cell_background(cell, bg_hex)
        set_cell_left_border(cell, color_hex, "24") # 3pt left border
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(f"⚠️ {title}: " if title == "IMPORTANT" else f"ℹ️ {title}: ")
        run_title.font.name = 'Calibri'
        run_title.font.size = Pt(11)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(14, 116, 144) if color_hex == "0E7490" else RGBColor(185, 28, 28)
        
        run_text = p.add_run(text)
        run_text.font.name = 'Calibri'
        run_text.font.size = Pt(11)
        run_text.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. Executive Architecture
    add_heading_1("1. Laboratory Overview & Architecture")
    
    p = doc.add_paragraph()
    p.add_run("In modern DevSecOps deployment structures, container workloads frequently share a singular base kernel infrastructure. If container controls are weakly configured, a local vulnerability can expand into a total host infrastructure compromise. This Hexa Force Laboratory demonstrates the four primary dimensions of container security isolation: Kernel separation, Namespace boundaries, Daemon interface control, and Volume access rules.")
    
    p = doc.add_paragraph()
    p.add_run("The laboratory provides interactive exploitation, detection system triggers, and mitigation automation to show how a pipeline secures itself at runtime. The workflow builds a single test image and executes the four stages sequentially, checking if defensive measures successfully prevent escape vectors.")
    
    # Table of Stages
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_widths = [Inches(1.2), Inches(1.8), Inches(1.8), Inches(1.7)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    # Style Header Row
    hdr_cells = table.rows[0].cells
    headers = ["SECURITY PILLAR", "STAGE FOCUS", "EXPLOIT SIMULATION", "DEFENSIVE KERNEL CONTROL"]
    for i, title in enumerate(headers):
        set_cell_background(hdr_cells[i], "0E7490")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(title)
        run.font.name = 'Trebuchet MS'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    stage_data = [
        ("Host Kernel Isolation", "Dirty COW (CVE-2016-5195) Race Condition", "Write race threads targeting a read-only root file", "Upgraded host kernels & Sandboxing runtimes (gVisor)"),
        ("Process Boundaries", "PID Namespace & Capability Sets", "Abusing CAP_SYS_PTRACE capabilities & Host PID sharing", "Dropping default permissions & Dropping CAP_SYS_PTRACE"),
        ("Daemon & API Security", "Exposed docker.sock UNIX socket", "HTTP POST to control socket & Spawn privilege containers", "Disabling daemon mounts & rootless runtime engines"),
        ("Volume Security", "Writable Path Traversal Mount", "Host Cron directory traversal write injection", "Mounting host volumes as read-only (:ro) configurations")
    ]
    
    for row_idx, data in enumerate(stage_data):
        row_cells = table.rows[row_idx + 1].cells
        # Apply zebra striping
        bg_hex = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate(data):
            set_cell_background(row_cells[i], bg_hex)
            set_cell_margins(row_cells[i], top=80, bottom=80, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(14, 116, 144)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 2. Detailed Stages
    add_heading_1("2. Modular Technical Breakdowns")

    # Stage 1
    add_heading_2("Stage 1: Host Kernel Isolation (Dirty COW - CVE-2016-5195)")
    p = doc.add_paragraph()
    p.add_run("Kernel-level sharing is the core efficiency of container structures. In this stage, an unprivileged user process inside the container runs threads competing to exploit the Dirty COW race condition. One thread calls ")
    p.add_run("madvise(..., MADV_DONTNEED)").font.bold = True
    p.add_run(" on a privately mapped read-only file, while another concurrently writes to ")
    p.add_run("/proc/self/mem").font.bold = True
    p.add_run(". This race exploits a non-atomic copy-on-write page fault lookup sequence to corrupt the read-only file in the shared page cache.")
    
    add_callout(
        "Under modern patched host kernels (such as the standard GitHub Actions host runners), this exploit is automatically serialization-blocked. The page tables serialize the writes safely, preserving target integrity.",
        "EXPLOIT NOTE", "0E7490", "F0F9FF"
    )

    # Stage 2
    add_heading_2("Stage 2: Process & Capabilities Isolation (CAP_SYS_PTRACE)")
    p = doc.add_paragraph()
    p.add_run("This stage evaluates boundary enforcement at the process level. If host process sharing is enabled (")
    p.add_run("--pid=host").font.bold = True
    p.add_run(") and combined with over-privileged process debugging capabilities (")
    p.add_run("CAP_SYS_PTRACE").font.bold = True
    p.add_run("), the container process can attach to and inject shellcode into host-level services, escaping the sandbox.")
    
    p = doc.add_paragraph()
    p.add_run("The pipeline actively runs process namespace isolation diagnostics using ")
    p.add_run("capsh --print").font.bold = True
    p.add_run(" and evaluating active namespace process PID counts. If the process PID counts are low, process boundary isolation is verified as active.")

    # Stage 3
    add_heading_2("Stage 3: Daemon & API Security (Exposed docker.sock)")
    p = doc.add_paragraph()
    p.add_run("The Docker daemon socket (")
    p.add_run("/var/run/docker.sock").font.bold = True
    p.add_run(") handles API requests as root. Mounting this socket inside containers is highly dangerous, as it grants full control over the engine. In this stage, the orchestrator detects the socket interface and models how a compromised user executes curl commands to trigger high-privilege sibling containers containing host root mounts.")
    
    p = doc.add_paragraph()
    p.add_run("The script successfully identifies the exposed socket interface and demonstrates the exact HTTP API payload signature required to execute the breakout, outlining why socket exposure is equivalent to giving root host permissions.")

    # Stage 4
    add_heading_2("Stage 4: Volume & Filesystem Security (Path Traversal)")
    p = doc.add_paragraph()
    p.add_run("Mounting sensitive host directories gives containers local system access. Stage 4 checks whether critical directories (simulated via ")
    p.add_run("/mnt/host/etc/cron.d").font.bold = True
    p.add_run(") are writable by the unprivileged container user. A breakout attempt is simulated by attempting a cron-job script injection.")
    
    p = doc.add_paragraph()
    p.add_run("In our secured environment, the directory is protected and owned by root. The script safely catches the ")
    p.add_run("Permission Denied").font.bold = True
    p.add_run(" output, logs the event, and prints a success notification indicating that the read-only file permissions correctly stopped the injection vector.")

    # 3. Pipeline Integration
    add_heading_1("3. Pipeline Integration & Automation")
    
    p = doc.add_paragraph()
    p.add_run("The entire lab suite is automated using a single GitHub Actions workflow file. The container is compiled from a hardened ")
    p.add_run("Dockerfile").font.bold = True
    p.add_run(" and executed using the modular shell orchestrator ")
    p.add_run("run_demo.sh").font.bold = True
    p.add_run(". Below is the complete, validated YAML structure of the CI/CD pipeline:")

    # GHA YAML block
    add_code_block([
        "name: Dirty COW Lab CI Pipeline",
        "",
        "on:",
        "  push:",
        "    branches: [ main ]",
        "  pull_request:",
        "    branches: [ main ]",
        "  workflow_dispatch: # Allows manual execution",
        "",
        "jobs:",
        "  build-and-run:",
        "    runs-on: ubuntu-latest",
        "    permissions:",
        "      contents: read",
        "",
        "    steps:",
        "    - name: Checkout Repository",
        "      uses: actions/checkout@v4",
        "      with:",
        "        token: ${{ secrets.GH_PAT }}",
        "",
        "    - name: Build Docker Container",
        "      run: |",
        "        echo \"=== Building the Dirty COW Lab Docker Container ===\"",
        "        docker build -t dirtycow-lab .",
        "",
        "    - name: \"Run Stage 1: Host Kernel Isolation (Dirty COW Exploit)\"",
        "      run: |",
        "        docker run --rm dirtycow-lab /bin/bash run_demo.sh --stage-1",
        "",
        "    - name: \"Run Stage 2: Namespace & Capabilities Isolation\"",
        "      run: |",
        "        docker run --rm dirtycow-lab /bin/bash run_demo.sh --stage-2",
        "",
        "    - name: \"Run Stage 3: Daemon & API Security (docker.sock Mount)\"",
        "      run: |",
        "        docker run --rm dirtycow-lab /bin/bash run_demo.sh --stage-3",
        "",
        "    - name: \"Run Stage 4: Volume & Filesystem Security (Path Traversal)\"",
        "      run: |",
        "        docker run --rm dirtycow-lab /bin/bash run_demo.sh --stage-4"
    ])

    # 4. Mitigation Checklist
    add_heading_1("4. DevSecOps Hardening & Remediation Checklist")
    
    p = doc.add_paragraph()
    p.add_run("To secure production environments and permanently stop these container escape vectors, execute the following remediation rules across your infrastructure:")

    # Checklist
    checklist_points = [
        ("Host Kernel Patching", "Update host Linux kernels immediately to ensure page fault lookups are atomic, eliminating Dirty COW race vectors."),
        ("Enforce Read-Only Mounts", "Always mount host volumes in containers using the read-only flag. Example: docker run -v /etc/config:/config:ro ubuntu."),
        ("Drop Default Capabilities", "Execute containers with stripped capability lists, explicitly dropping CAP_SYS_PTRACE and CAP_SYS_ADMIN capabilities."),
        ("Configure Low-Privilege Users", "Switch away from the root user inside images by declaring USER victim in Dockerfiles to prevent root filesystem writes."),
        ("Eliminate Socket Mounting", "Never mount /var/run/docker.sock into untrusted pods. Migrate workloads to secure rootless runtime engines (e.g. Podman).")
    ]

    for title, desc in checklist_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(f"[{title}] ")
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(14, 116, 144)
        
        run_desc = p.add_run(desc)

    # Save
    output_path = "Hexa_Force_Implementation_Guide.docx"
    doc.save(output_path)
    print(f"[+] Document successfully generated at: {os.path.abspath(output_path)}")

if __name__ == '__main__':
    create_document()
