import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="8" w:space="0" w:color="888888"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def make_callout(doc, text, title="PRO-TIP FOR PRESENTING"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=180, bottom=180, left=240, right=240)
    
    # Left border thick blue
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2563EB"/><w:top w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"★ {title}\n")
    run_title.font.name = 'Segoe UI'
    run_title.font.bold = True
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(37, 99, 235)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Segoe UI'
    run_text.font.italic = True
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(51, 65, 85)

def main():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Segoe UI'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)
    
    h1 = styles['Heading 1']
    h1.font.name = 'Segoe UI'
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    
    h2 = styles['Heading 2']
    h2.font.name = 'Segoe UI'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(37, 99, 235) # Blue 600
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # 1. COVER PAGE
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("HEXA FORCE")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 41, 59)
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    sub_run = sub_p.add_run("Presentation Companion & Speaker Talk-Tracks\nContainer Security Attack Lab & Dirty COW Deep Dive")
    sub_run.font.name = 'Segoe UI'
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(71, 85, 105)
    
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(200)
    meta_run = meta_p.add_run("Prepared by: Hexa Force Core Engineering Team\nPurpose: Complete Boardroom Delivery & Examiner Q&A Guide\nVersion: 2.0 (Hardened)\nStatus: Production Ready")
    meta_run.font.name = 'Segoe UI'
    meta_run.font.size = Pt(11)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()
    
    # 2. INTRODUCTION SECTION
    doc.add_heading("How to Use This Speaker Guide", level=1)
    intro_p = doc.add_paragraph(
        "Welcome to the Hexa Force Presentation Companion. This guide has been compiled to act as your "
        "ultimate reference while showing the PowerPoint presentation decks to examiners or corporate "
        "stakeholders. It covers two presentation decks: (1) The Hexa Force Secure Lab Presentation, and (2) "
        "The Dirty COW Technical Briefing. "
    )
    intro_p.add_run(
        "\n\nEach slide is broken down with direct verbatim talk tracks, a list of critical technical keywords to hit "
        "so you sound authoritative, and pro-active Q&A entries showing tough questions the evaluators are "
        "highly likely to ask you, along with the precise answers you should give to secure maximum points."
    )
    
    make_callout(
        doc,
        "Keep this document open on a side monitor, a tablet, or printed in front of you. Speak slowly, project "
        "confidence, and lean heavily on the precise technical terms (like namespaces, cgroups, copy-on-write, and "
        "OIDC federation) detailed in this guide.",
        "EXAMINER PREPARATION STRATEGY"
    )
    
    doc.add_page_break()
    
    # 3. DECK 1: HEXA FORCE SECURE LAB PRESENTATION
    doc.add_heading("Deck 1: Hexa Force Secure Lab Presentation", level=1)
    doc.add_paragraph(
        "This deck walks through the complete 4-stage container attack and hardening lab. It shows "
        "how a minor privilege boundary lapse can lead to physical host takeover, and how modern cryptographic supply "
        "chain signing with Cosign and OIDC neutralizes container image manipulation."
    )
    
    slides_deck1 = [
        {
            "num": 1,
            "title": "Title Slide: Hexa Force Container Security Lab",
            "summary": "Title and introduction layout displaying team branding, lab title, and speaker information.",
            "verbatim": (
                "\"Good morning, evaluators and members of the panel. Welcome to the technical delivery of the "
                "Hexa Force Container Security Lab. Today, we are presenting our end-to-end container security "
                "demonstration, modeling a high-fidelity 4-stage attack chain and corresponding defense-in-depth "
                "hardening strategies. Over the next few minutes, we will walk you through how a local privilege "
                "escalation vulnerability can cascade into complete host takeover, and how modern cryptographic verification "
                "using keyless Sigstore/Cosign signing with GitHub Actions OIDC robustly secures the software delivery pipeline. "
                "Let us begin by reviewing our core objectives.\""
            ),
            "keywords": ["Hexa Force Security Lab", "4-Stage Attack Chain", "Cryptographic supply chain verification", "Keyless signing", "Defense-in-Depth"],
            "q": "What motivated the choice of this specific attack chain?",
            "a": "We chose this specific progression because it realistically models how simple runtime misconfigurations and unpatched host-level vulnerabilities interact in enterprise microservices. It demonstrates that container security cannot rely on runtime isolation alone, but must encompass host hardening, image provenance, and CI/CD integrity."
        },
        {
            "num": 2,
            "title": "Executive Summary & Lab Objectives",
            "summary": "Split-column showing the dual mission: simulating container vulnerabilities and enforcing production-grade security remediations.",
            "verbatim": (
                "\"Our mission at Hexa Force was two-fold. First, to construct a highly visual, realistic container exploit "
                "environment to study how isolation boundaries fail; and second, to design and execute production-ready "
                "remediations. We prove that security should not be a runtime barrier, but a foundational, automated system. "
                "We will walk you through a sequence where we compromise a container, abuse Namespace/Capability settings, "
                "access the Docker API socket, and perform a full host filesystem breakout. We then demonstrate how to shut "
                "this chain down completely using capability restriction, read-only structures, and Sigstore verification.\""
            ),
            "keywords": ["Boundary failure replication", "Isolation boundaries", "Capability restriction", "Sigstore container verification"],
            "q": "Why is privilege escalation inside a container dangerous if containers are isolated?",
            "a": "Container isolation relies on kernel namespaces and cgroups. If an attacker elevates to 'root' inside a container, namespaces still limit what they can see. However, if the container is run with excessive Linux capabilities (like CAP_SYS_ADMIN) or has access to host sockets, a container root user can easily leverage these privileges to break out of the isolation entirely."
        },
        {
            "num": 3,
            "title": "Lab Architecture & Threat Model",
            "summary": "Technical diagram showing a low-privileged container running under the 'victim' user, sharing the host kernel directly.",
            "verbatim": (
                "\"Let's establish the architectural baseline. The entire lab runs inside an Ubuntu-based Docker container "
                "that shares the host's Linux kernel directly. The application inside the container runs as a low-privileged "
                "non-root user named 'victim'. This represents a realistic threat model, simulating an external application compromise "
                "where an attacker secures a standard shell. The underlying host kernel is intentionally left unpatched against "
                "CVE-2016-5195 to demonstrate the catastrophic vulnerability inherent to shared kernel microservice environments.\""
            ),
            "keywords": ["Unprivileged container user 'victim'", "Shared Linux kernel", "Threat modeling", "CVE-2016-5195 exposure"],
            "q": "Do virtual machines (VMs) share this same kernel vulnerability?",
            "a": "No. Virtual machines run on a hypervisor with their own dedicated guest operating system and guest kernel, isolating them at the hardware level. Standard Docker containers, however, share the host kernel directly. This is a critical distinction: if the host kernel is vulnerable, all standard containers running on it are potentially exposed."
        },
        {
            "num": 4,
            "title": "Stage 1: Dirty COW Local Privilege Escalation",
            "summary": "Terminal view showing the compilation and execution of the dirtyc0w.c exploit to bypass read-only files.",
            "verbatim": (
                "\"At Stage 1, we execute our first exploit. Running as the unprivileged 'victim' user, we execute our compiled "
                "binary, 'dirtyc0w'. This exploit targets the Dirty COW kernel vulnerability, which is a race condition in the "
                "kernel's virtual memory subsystem's handling of Copy-on-Write. We target a read-only secret file, '/var/secret/target.txt'. "
                "By using madvise with the MADV_DONTNEED flag, we race the kernel's write mechanisms, tricking the kernel into "
                "writing root-owned data into a read-only memory segment. This instantly elevates our privilege from unprivileged user to container root.\""
            ),
            "keywords": ["Virtual memory race condition", "Copy-on-Write (COW)", "madvise(MADV_DONTNEED)", "Container-level root elevation"],
            "q": "What exactly does the 'madvise' system call do in this exploit?",
            "a": "The madvise system call with the MADV_DONTNEED flag advises the kernel that the calling application no longer expects to access the specified memory range. The kernel responds by freeing the physical memory page table entries, causing a subsequent read to load the data from the read-only backing file again, creating a race window for the write process to write to physical memory before the new COW copy is allocated."
        },
        {
            "num": 5,
            "title": "Stage 2: Namespace Isolation & Capability Abuse",
            "summary": "Listing of Linux capabilities, showing how the presence of CAP_SYS_ADMIN breaks namespace boundaries.",
            "verbatim": (
                "\"In Stage 2, we exploit excessive container privileges. Having achieved container root in Stage 1, we now check "
                "the container's capabilities. If a container is run with elevated Linux capabilities, such as CAP_SYS_ADMIN, "
                "the boundaries of standard container isolation begin to disintegrate. Linux capabilities partition traditional "
                "superuser root powers into distinct, assignable units. By abusing CAP_SYS_ADMIN, we gain the authority to modify "
                "kernel namespaces and mount external devices, laying the physical groundwork for our escape to the host filesystem.\""
            ),
            "keywords": ["CAP_SYS_ADMIN capability", "Boundary disintegration", "Superuser privilege partitioning", "Namespace modification permissions"],
            "q": "Why would an administrator ever enable CAP_SYS_ADMIN in production?",
            "a": "Administrators often enable CAP_SYS_ADMIN or grant the full '--privileged' flag to support legacy monitoring agents, VPN clients, backup utilities, or nested Docker builds. It highlights the absolute necessity of performing strict capability profiling and denying administrative privileges by default."
        },
        {
            "num": 6,
            "title": "Stage 3: Shared Docker Socket & Container Escape",
            "summary": "Visual diagram showing communication from inside the container to the host's Docker socket (/var/run/docker.sock).",
            "verbatim": (
                "\"Stage 3 demonstrates a critical lateral movement vector: the exposed Docker socket. In many development "
                "environments, administrators mount the host's Docker socket—'/var/run/docker.sock'—directly into the container. "
                "Since our Stage 1 exploit gave us root privileges inside the container, we have full read/write permission to "
                "this socket. Because the Docker socket is the direct command API to the host's Docker engine, we can now send API requests "
                "directly to the host daemon. This effectively grants us full control over the host's Docker operations.\""
            ),
            "keywords": ["/var/run/docker.sock exposure", "Docker API communications", "Host Docker daemon access", "API-driven container creation"],
            "q": "How can you securely isolate containers if they need to build Docker images?",
            "a": "Instead of exposing the host's Docker socket, you should use daemonless, unprivileged container build engines such as Kaniko or Podman. They do not require a running Docker socket to build images and compile layers, preserving the isolation boundary completely."
        },
        {
            "num": 7,
            "title": "Stage 4: Volume Mount Exploitation & Host Cron Takeover",
            "summary": "Visual progression of launching a new privileged container, mounting the host filesystem (/), and injecting a persistent cron reverse shell.",
            "verbatim": (
                "\"In Stage 4, we execute the final breakout. Using our access to the Docker socket, we issue an API command to "
                "launch a new, highly privileged container on the host. In its configuration, we mount the host's physical root filesystem "
                "to '/mnt/host' inside the container. Because this container runs with full host namespace access, we traverse "
                "directly into '/mnt/host/etc/cron.d' and write a malicious cron job. When the host's cron daemon executes, "
                "it runs our script and opens a root-privileged reverse shell back to our attack machine. We have achieved complete host takeover.\""
            ),
            "keywords": ["Host root filesystem mount (/)", "/etc/cron.d cron injection", "Root reverse shell execution", "Complete physical host takeover"],
            "q": "Why is injecting a cron job better for an attacker than just running a command directly on the host?",
            "a": "Injecting a cron job establishes persistent access. If the attacker simply runs a single process, it might be terminated, or the host might reboot. Writing to `/etc/cron.d` ensures that even after a system restart, the host daemon will periodically reconnect the reverse shell, giving the attacker long-term control."
        },
        {
            "num": 8,
            "title": "Defense-in-Depth: Hardening & Remediation Strategies",
            "summary": "Comparison table displaying 'Default Vulnerable Configuration' vs 'Hexa Force Hardened Configuration'.",
            "verbatim": (
                "\"Now, let us examine how we completely remediate this threat. Hexa Force implements a comprehensive "
                "Defense-in-Depth architecture. First, we restrict container capabilities by explicitly dropping CAP_SYS_ADMIN. "
                "Second, we run the container as a non-root user and mark the root filesystem as read-only. Third, we block "
                "all exposure of the Docker socket. Finally, we apply kernel security patches to eliminate Dirty COW at the host level. "
                "Even if an attacker gains an initial shell, these combined controls prevent them from ever gaining root or escaping the container.\""
            ),
            "keywords": ["Defense-in-Depth architecture", "Capability dropping", "Read-only container rootfs", "Host-level kernel security patching"],
            "q": "What is the security value of a read-only container root filesystem?",
            "a": "A read-only root filesystem prevents an attacker from writing files, installing tools, downloading exploit scripts, or compiling C binaries (like `dirtyc0w.c`) inside the container. It dramatically reduces the available attack surface and prevents runtime environmental tampering."
        },
        {
            "num": 9,
            "title": "Cryptographic Keyless Container Signing",
            "summary": "Supply chain diagram showing GitHub Actions OIDC identity token exchange and keyless Cosign signature recording in the Rekor public ledger.",
            "verbatim": (
                "\"Beyond runtime hardening, Hexa Force secures the software supply chain using state-of-the-art keyless "
                "container signing via Cosign. Traditional signing relies on storing highly sensitive cryptographic private keys. "
                "If these keys are leaked, the supply chain is compromised. Hexa Force avoids this completely. We leverage OpenID Connect "
                "to exchange GitHub Action build identities for temporary, 10-minute certificates. Cosign signs the built image, "
                "recording the signature in the immutable public Rekor transparency log, guaranteeing container origin and integrity.\""
            ),
            "keywords": ["Keyless Cosign container signing", "OpenID Connect (OIDC) identity federation", "Rekor public transparency log", "Software supply chain integrity"],
            "q": "If there are no private keys stored, how is the signature verified later?",
            "a": "Verification is done by checking the short-lived certificate associated with the signature. The signature is bound to a specific OIDC identity (like the GitHub repo URL and workflow path). During verification, Cosign confirms that the public certificate was signed by the Sigstore Fulcio Certificate Authority during its valid 10-minute window, and that the transaction was logged in the tamper-proof Rekor ledger."
        },
        {
            "num": 10,
            "title": "CI/CD Pipeline Automation (GitHub Actions & GHCR)",
            "summary": "GitHub Actions workflow structure detailing automated building, tag lowercasing, GHCR pushing, and OIDC container signing.",
            "verbatim": (
                "\"We operationalize this supply chain defense by automating the entire pipeline inside GitHub Actions. Upon a push "
                "to the main branch, our workflow triggers. It builds the Docker image, dynamically converts the image tag to "
                "lowercase using bash expansion to ensure compatibility with container registry specifications, and pushes it to "
                "the GitHub Container Registry. Immediately, Cosign uses the OIDC identity to sign the image. This seamless, "
                "automated pipeline guarantees that only verified and signed code is allowed to run in our production environment.\""
            ),
            "keywords": ["GitHub Actions automation", "GitHub Container Registry (GHCR)", "Image tag lowercase conversion", "Automated deployment gates"],
            "q": "Why is the lowercase conversion step necessary in the GHA workflow?",
            "a": "The Docker and OCI image specification requires all image names and registry repository paths to be strictly lowercase. Since GitHub usernames and repository names can contain uppercase letters (like `AzDevops143`), we convert the string dynamically using `${GITHUB_REPOSITORY,,}` in the shell to prevent build failures during push."
        },
        {
            "num": 11,
            "title": "Summary of Results & Q&A Preparation",
            "summary": "Conclusion layout with project achievements and a welcoming prompt for examiner questions.",
            "verbatim": (
                "\"In conclusion, the Hexa Force Container Security Lab successfully demonstrates the extreme vulnerability "
                "of default container runtimes, while delivering an enterprise-ready blueprint for comprehensive defense. "
                "By pairing strict runtime hardening policies with cryptographic supply chain verification, we ensure complete "
                "resilience against host compromises and supply chain injection. Thank you for your time. The Hexa Force team is "
                "now open to your technical questions.\""
            ),
            "keywords": ["Vulnerability mitigation verification", "Enterprise hardening blueprint", "Supply chain injection resilience", "Open floor Q&A session"],
            "q": "What is the most critical takeaway from your lab?",
            "a": "The most critical takeaway is that container security is a holistic discipline. Isolation is not a static lock; it requires continuous kernel maintenance, strict adherence to the principle of least privilege at runtime, and supply-chain verification through modern cryptographic signing."
        }
    ]

    for slide in slides_deck1:
        doc.add_heading(f"Slide {slide['num']}: {slide['title']}", level=2)
        
        # Details Table
        table = doc.add_table(rows=4, cols=2)
        set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        col_widths = [Inches(1.8), Inches(4.7)]
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
        
        # Row 1: Visual Summary
        cell_l1 = table.cell(0, 0)
        cell_r1 = table.cell(0, 1)
        set_cell_background(cell_l1, "F8FAFC")
        set_cell_margins(cell_l1)
        set_cell_margins(cell_r1)
        run_l1 = cell_l1.paragraphs[0].add_run("Slide Visuals:")
        run_l1.bold = True
        cell_r1.paragraphs[0].text = slide["summary"]
        
        # Row 2: Verbatim Talk Track
        cell_l2 = table.cell(1, 0)
        cell_r2 = table.cell(1, 1)
        set_cell_background(cell_l2, "F8FAFC")
        set_cell_margins(cell_l2)
        set_cell_margins(cell_r2)
        run_l2 = cell_l2.paragraphs[0].add_run("Spoken Talk Track\n(Verbatim Script):")
        run_l2.bold = True
        cell_r2.paragraphs[0].text = slide["verbatim"]
        cell_r2.paragraphs[0].runs[0].font.italic = True
        
        # Row 3: Key Keywords
        cell_l3 = table.cell(2, 0)
        cell_r3 = table.cell(2, 1)
        set_cell_background(cell_l3, "F8FAFC")
        set_cell_margins(cell_l3)
        set_cell_margins(cell_r3)
        run_l3 = cell_l3.paragraphs[0].add_run("Key Terms to Hit:")
        run_l3.bold = True
        
        # Bullet list for keywords
        kw_p = cell_r3.paragraphs[0]
        kw_p.text = ""
        for kw in slide["keywords"]:
            r = kw_p.add_run(f"• {kw}   ")
            r.bold = True
            r.font.color.rgb = RGBColor(37, 99, 235)
            
        # Row 4: Anticipated Q&A
        cell_l4 = table.cell(3, 0)
        cell_r4 = table.cell(3, 1)
        set_cell_background(cell_l4, "F8FAFC")
        set_cell_margins(cell_l4)
        set_cell_margins(cell_r4)
        run_l4 = cell_l4.paragraphs[0].add_run("Examiner Q&A Prep:")
        run_l4.bold = True
        
        qa_p = cell_r4.paragraphs[0]
        qa_p.text = ""
        q_run = qa_p.add_run(f"Question: {slide['q']}\n")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(15, 23, 42)
        a_run = qa_p.add_run(f"Strategic Answer: {slide['a']}")
        a_run.font.color.rgb = RGBColor(71, 85, 105)
        
        doc.add_paragraph().paragraph_format.space_before = Pt(6)
        make_callout(doc, f"Make sure to emphasize: {slide['keywords'][0]} and speak confidently during this slide.", f"SPEAKER TIP FOR SLIDE {slide['num']}")
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # 4. DECK 2: DIRTY COW TECHNICAL BRIEFING
    doc.add_heading("Deck 2: Dirty COW Technical Briefing", level=1)
    doc.add_paragraph(
        "This deck provides a deep dive on CVE-2016-5195 (Dirty COW), examining Linux memory "
        "management, copy-on-write page cache operations, race conditions, and kernel mitigation mechanisms."
    )
    
    slides_deck2 = [
        {
            "num": 1,
            "title": "Title & Technical Scope",
            "summary": "Technical title slide with deep-dive focus: kernel virtual memory exploits and Copy-on-Write race conditions.",
            "verbatim": (
                "\"Welcome to Part 2 of our briefing: The Dirty COW Technical Deep Dive. In this session, we will demystify "
                "the precise mechanisms behind CVE-2016-5195, commonly known as Dirty COW. This kernel exploit bypassed "
                "decades of security policies in Linux systems by weaponizing a race condition inside the virtual memory management "
                "subsystem. We will explore how memory pages are mapped, how Copy-on-Write page cache mechanisms work, and "
                "how we can force the kernel to write user data directly to read-only root files. Let us dive into the exploit's background.\""
            ),
            "keywords": ["CVE-2016-5195 Technical Deep Dive", "Virtual memory subsystem exploitation", "Copy-on-Write core vulnerability", "Kernel boundary bypass"],
            "q": "Why was the name 'Dirty COW' chosen?",
            "a": "The name is a technical pun. 'COW' stands for the 'Copy-on-Write' mechanism in operating system memory management, and 'Dirty' refers to marking a cached memory page as modified (or 'dirty') in the page table, indicating it needs to be written back to disk."
        },
        {
            "num": 2,
            "title": "Exploit Context & CVE-2016-5195 Profile",
            "summary": "Vulnerability database summary displaying CVSS 3.x scores, affected kernel versions, and attack vector details.",
            "verbatim": (
                "\"To contextualize this threat, let's examine the profile of CVE-2016-5195. Discovered in 2016, this vulnerability "
                "carries a High CVSS score. It affected almost every Linux kernel version prior to October 2016. The beauty—and "
                "danger—of Dirty COW is that it requires absolutely no special system privileges to execute. Any local user, even "
                "an unprivileged process running inside a Docker container, can trigger this exploit to bypass file permissions "
                "and write root-owned configurations. Let us look at the memory architecture that made this possible.\""
            ),
            "keywords": ["CVSS high-severity rating", "Universal local exploit vector", "Kernel version scope", "File permission bypass"],
            "q": "Can Dirty COW be exploited remotely?",
            "a": "No. Dirty COW is strictly a local privilege escalation (LPE) vulnerability. The attacker must already have local execution capabilities—such as an unprivileged shell inside a container or a standard system terminal—to execute the binary and exploit the kernel."
        },
        {
            "num": 3,
            "title": "Virtual Memory Management & COW Mechanics",
            "summary": "Logical schematic of Linux page mapping, showing private page mapping and Copy-on-Write allocation flow.",
            "verbatim": (
                "\"To understand the exploit, we must first understand virtual memory and Copy-on-Write. In Linux, when a process "
                "requests a private copy of a file or memory mapping, the kernel maps the process's virtual memory address to the "
                "same physical memory page as the source file to conserve RAM. This is called a shared private mapping. "
                "Under normal operations, if the process attempts to write to this page, the kernel detects it, allocates a new "
                "physical page, copies the original content, and updates the page tables for the writing process. This is the Copy-on-Write mechanism.\""
            ),
            "keywords": ["Virtual memory address space", "Shared private page mappings", "Copy-on-Write RAM optimization", "Page table allocation"],
            "q": "What is the difference between a shared mapping and a private mapping in Linux memory?",
            "a": "A shared mapping (`MAP_SHARED`) means that any write to the memory is written back to the underlying backing file directly. A private mapping (`MAP_PRIVATE`) means that writes are kept local to the process using Copy-on-Write, ensuring the backing file remains completely unmodified."
        },
        {
            "num": 4,
            "title": "The Vulnerability: Race Condition in madvise & write",
            "summary": "Concurrent thread race flowchart showcasing Thread A calling write and Thread B calling madvise(MADV_DONTNEED).",
            "verbatim": (
                "\"The vulnerability is a race condition. When a process tries to write to a private read-only mapping, the kernel "
                "initiates the COW process. First, it locates the physical page, then allocates a new page, and finally updates "
                "the page tables. However, this multi-step process is not atomic. In a multi-threaded system, Thread A constantly "
                "attempts to write to the memory. At the exact same time, Thread B calls the 'madvise' system call with "
                "'MADV_DONTNEED'. This advises the kernel that the mapping is no longer needed, forcing the kernel to discard the "
                "new page. The write operation races this, writing directly to the underlying read-only backing file before the new page table entry is bound.\""
            ),
            "keywords": ["Non-atomic kernel operations", "Thread race condition concurrency", "madvise(MADV_DONTNEED) page disposal", "Read-only backing file corruption"],
            "q": "Why is the race condition so highly successful in elevating privileges?",
            "a": "Because the exploit runs the write system call and the madvise system call in a high-speed loop across multiple CPU cores. By executing millions of times per second, the exploit guarantees that the kernel will eventually execute the write request during the exact microsecond window after the memory page is discarded but before the new page is allocated."
        },
        {
            "num": 5,
            "title": "Execution Flowchart & Code Walkthrough (dirtyc0w.c)",
            "summary": "Source code display of the two threads in dirtyc0w.c, highlighting mmap, madvise, and write loops.",
            "verbatim": (
                "\"Let us analyze the actual exploit code, 'dirtyc0w.c'. The program begins by opening the target file in read-only "
                "mode and mapping it using mmap with 'MAP_PRIVATE'. It then spawns two parallel threads. Thread 1 runs a loop calling "
                "madvise with the MADV_DONTNEED flag on the mapped memory range. Thread 2 opens '/proc/self/mem'—which is the interface "
                "to the process's own memory—and continuously writes to the mapped address. By abusing the /proc/self/mem interface, "
                "the exploit bypasses standard write permission checks, forcing the kernel to execute the write against the physical read-only page.\""
            ),
            "keywords": ["mmap with MAP_PRIVATE", "/proc/self/mem write interface", "Dual-threaded exploit loops", "Permission check bypass"],
            "q": "Why does the exploit use '/proc/self/mem' instead of a normal pointer write?",
            "a": "Writing directly via a pointer would cause the compiler or the CPU MMU to immediately throw a segmentation fault, because the memory is mapped read-only. By writing to `/proc/self/mem`, the write request goes through the kernel's virtual memory system interface, allowing us to exploit the race condition within the kernel itself."
        },
        {
            "num": 6,
            "title": "Remediation, Patch Analysis & Enterprise Mitigation",
            "summary": "Kernel patch diff display showing the introduction of the 'FOLL_COW' flag inside virtual memory handlers.",
            "verbatim": (
                "\"Finally, let us review the patch. The Linux kernel community resolved this vulnerability by introducing a new "
                "internal flag called 'FOLL_COW' in the memory handling subsystem. The patch ensures that when the kernel performs "
                "a write follow operation, it checks if a Copy-on-Write operation has already occurred and preserves this state even "
                "if madvise is called concurrently. For enterprise environments, the core remediation is simple: keep the host kernel "
                "updated. In cloud-native systems, always verify that your host operating system is fully patched to secure "
                "your container runtime environment.\""
            ),
            "keywords": ["FOLL_COW patch implementation", "Write follow operation state preservation", "Host OS patching core mandate", "Container host kernel security"],
            "q": "What is FOLL_COW and how does it prevent the race?",
            "a": "Before the patch, if a page table entry was discarded by madvise, the kernel would forget that the write was originally a Copy-on-Write operation. The introduced `FOLL_COW` flag serves as an internal sticky bit, ensuring the kernel tracks that it is in the middle of a COW process, preventing it from falling back to the original read-only page."
        }
    ]

    for slide in slides_deck2:
        doc.add_heading(f"Slide {slide['num']}: {slide['title']}", level=2)
        
        # Details Table
        table = doc.add_table(rows=4, cols=2)
        set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        col_widths = [Inches(1.8), Inches(4.7)]
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
        
        # Row 1: Visual Summary
        cell_l1 = table.cell(0, 0)
        cell_r1 = table.cell(0, 1)
        set_cell_background(cell_l1, "F8FAFC")
        set_cell_margins(cell_l1)
        set_cell_margins(cell_r1)
        run_l1 = cell_l1.paragraphs[0].add_run("Slide Visuals:")
        run_l1.bold = True
        cell_r1.paragraphs[0].text = slide["summary"]
        
        # Row 2: Verbatim Talk Track
        cell_l2 = table.cell(1, 0)
        cell_r2 = table.cell(1, 1)
        set_cell_background(cell_l2, "F8FAFC")
        set_cell_margins(cell_l2)
        set_cell_margins(cell_r2)
        run_l2 = cell_l2.paragraphs[0].add_run("Spoken Talk Track\n(Verbatim Script):")
        run_l2.bold = True
        cell_r2.paragraphs[0].text = slide["verbatim"]
        cell_r2.paragraphs[0].runs[0].font.italic = True
        
        # Row 3: Key Keywords
        cell_l3 = table.cell(2, 0)
        cell_r3 = table.cell(2, 1)
        set_cell_background(cell_l3, "F8FAFC")
        set_cell_margins(cell_l3)
        set_cell_margins(cell_r3)
        run_l3 = cell_l3.paragraphs[0].add_run("Key Terms to Hit:")
        run_l3.bold = True
        
        # Bullet list for keywords
        kw_p = cell_r3.paragraphs[0]
        kw_p.text = ""
        for kw in slide["keywords"]:
            r = kw_p.add_run(f"• {kw}   ")
            r.bold = True
            r.font.color.rgb = RGBColor(37, 99, 235)
            
        # Row 4: Anticipated Q&A
        cell_l4 = table.cell(3, 0)
        cell_r4 = table.cell(3, 1)
        set_cell_background(cell_l4, "F8FAFC")
        set_cell_margins(cell_l4)
        set_cell_margins(cell_r4)
        run_l4 = cell_l4.paragraphs[0].add_run("Examiner Q&A Prep:")
        run_l4.bold = True
        
        qa_p = cell_r4.paragraphs[0]
        qa_p.text = ""
        q_run = qa_p.add_run(f"Question: {slide['q']}\n")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(15, 23, 42)
        a_run = qa_p.add_run(f"Strategic Answer: {slide['a']}")
        a_run.font.color.rgb = RGBColor(71, 85, 105)
        
        doc.add_paragraph().paragraph_format.space_before = Pt(6)
        make_callout(doc, f"Emphasize: {slide['keywords'][0]} - this is a high-difficulty academic concept. Clear delivery is key.", f"DEEP-DIVE TECHNICAL TIP")
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Save Document
    filename = "Hexa_Force_Presentation_Readout_Guide.docx"
    doc.save(filename)
    print(f"SUCCESS: Readout guide saved as {filename}")

if __name__ == "__main__":
    main()
